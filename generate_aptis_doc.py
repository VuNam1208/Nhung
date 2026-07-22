from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "/workspace/Aptis_Speaking_Part_4_8_Bai_Gop.docx"

LESSONS = [
    {
        "title": "Bài 1 — Chủ đề 01 + 02 + 13",
        "topics": (
            "A difficult question / job interview / challenge / achievement / "
            "learning English / learning a new skill / English course / great effort / "
            "good news / received a gift"
        ),
        "text": [
            "I would like to talk about a challenging experience that eventually became an important personal achievement.",
            "It began about a year ago when I decided to improve my English speaking skills. Although I could read and understand English reasonably well, I lacked confidence whenever I had to speak. Therefore, I joined an evening English course and practised for at least thirty minutes every day. Since I usually attended class after work, maintaining this routine required a great deal of effort.",
            "After several months, I applied for a part-time position at an international company. Before the interview, I had prepared answers about my education, strengths and previous experience, so I felt fairly confident. However, the manager suddenly asked me, “What is your biggest weakness, and how are you trying to improve it?”",
            "At that moment, my mind went blank because I did not want to make a bad impression. However, I took a deep breath and decided to be honest. I explained that I sometimes focused too much on minor details, which could slow me down. I also described how I was learning to set priorities and manage my time more effectively.",
            "A few days later, I was delighted to learn that I had been offered the job. My parents were extremely proud and gave me a watch to celebrate.",
            "Looking back, I realised that consistent effort and good preparation can help us overcome difficult challenges. Since then, I have become much more confident when speaking English and dealing with unexpected questions.",
        ],
        "focus": [
            "English course: nhấn mạnh quá trình học và luyện tập mỗi ngày.",
            "Interview/difficult question: nhấn mạnh câu hỏi khó và cách xử lý.",
            "Good news/gift/achievement: nhấn mạnh kết quả nhận việc và món quà.",
        ],
    },
    {
        "title": "Bài 2 — Chủ đề 03 + 04 + 06",
        "topics": (
            "A long trip / holiday / vacation / new city / place you travelled to / "
            "got lost / bad weather / in a hurry / visited a friend / met a new friend / "
            "were helped / laughed with a friend"
        ),
        "text": [
            "I would like to talk about a memorable trip I took to Da Nang to visit one of my closest friends.",
            "It happened last summer when I needed a break after a stressful period at work. Since my friend had recently moved to Da Nang, I decided to spend my holiday there. Before setting off, I had booked my train ticket, prepared my clothes and saved my friend’s address on my phone. As it was my first time visiting the city, I felt extremely excited.",
            "However, shortly after I arrived, the weather suddenly became worse and it started raining heavily. While trying to find my friend’s apartment, I took the wrong road and got completely lost. To make matters worse, my phone battery was almost dead, and I was in a hurry because my friend was waiting for me.",
            "At that moment, I felt worried and confused. However, instead of panicking, I asked a local student for help. He was incredibly friendly and showed me which bus to take. Thanks to his directions, I eventually reached my friend’s apartment safely.",
            "That evening, my friend introduced me to several of his new friends. We had dinner together and laughed about how I had managed to get lost despite using an online map.",
            "Looking back, the journey was stressful but unforgettable. It taught me that staying calm and asking for help can make difficult situations easier. Since then, I have always charged my phone, checked the weather forecast and saved important addresses before travelling.",
        ],
    },
    {
        "title": "Bài 3 — Chủ đề 05 + 09",
        "topics": (
            "Helped someone / teamwork / busy time / effort / activity for children / "
            "elderly people / volunteering"
        ),
        "text": [
            "I would like to talk about a time when I worked with my classmates to organise a volunteer activity.",
            "It happened last year when our university arranged a charity event at a local community centre. Our aim was to provide educational activities for children and spend time with elderly people who lived nearby. Before the event, we had prepared English games, small gifts and several group activities.",
            "At that time, everyone was extremely busy because we had only one week to complete the preparations. To make matters worse, one member of our team became ill and could not finish his work. Although I already had several responsibilities of my own, I decided to help him because his tasks were essential to the event.",
            "First, I helped organise his ideas and divided the work into smaller sections. I then prepared additional materials and asked the other team members to check them. We had to work late for several evenings, so it required a great deal of effort. However, nobody complained because we all wanted the event to succeed.",
            "In the end, the children thoroughly enjoyed the English games, while the elderly people were grateful for our visit. Although I felt exhausted afterwards, I was incredibly proud of what we had achieved together.",
            "Looking back, I realised that teamwork is not only about completing our own responsibilities but also about supporting others when they are struggling. Since then, I have become more willing to volunteer and more confident when working as part of a team.",
        ],
    },
    {
        "title": "Bài 4 — Chủ đề 07 + 08",
        "topics": (
            "Saved money / wanted to buy something but could not / many choices / "
            "planning something / many options / busy time / hurried"
        ),
        "text": [
            "I would like to talk about a time when I saved money and planned a surprise birthday gift for my younger brother.",
            "It happened about two years ago when his old laptop stopped working. Since he needed a computer for his studies, I wanted to buy him a new one for his birthday. However, the model he liked was quite expensive, and I could not afford it immediately.",
            "At first, I felt disappointed because I wanted to help him. In addition, there were so many different models that I found it difficult to make a decision. Some laptops were powerful but unaffordable, while cheaper ones did not have the features he needed. Instead of buying one impulsively, I decided to make a careful plan.",
            "For several months, I reduced unnecessary spending and put part of my salary aside. I also compared prices, watched product reviews and asked a knowledgeable friend for advice. Eventually, I found a laptop that was both reliable and reasonably priced.",
            "On my brother’s birthday, I was extremely busy. I had to collect the laptop after work, buy a cake and arrive home before he did. Although I was in a hurry, everything went according to plan. He was completely surprised and delighted with the gift.",
            "Looking back, I realised that having many choices does not necessarily make decisions easier. The experience taught me that patience, careful planning and sensible money management are essential. Since then, I have tried to avoid making impulsive purchases.",
        ],
    },
    {
        "title": "Bài 5 — Chủ đề 10",
        "topics": (
            "Sports match / music festival / amusement park / funny moment with friends"
        ),
        "text": [
            "I would like to talk about a memorable day I spent with my friends at a large entertainment festival.",
            "It happened last summer at an amusement park near my hometown. The event included live music, several outdoor games and a friendly football match between two local teams. One of my friends had received some free tickets, so he invited the rest of us to go with him.",
            "Before the event, I had not expected it to be particularly exciting because I was neither a huge football fan nor someone who regularly attended music festivals. However, when we arrived, I was immediately impressed by the atmosphere. Hundreds of people were cheering, singing and taking part in different activities.",
            "We watched the football match first. The funniest moment occurred when one of my friends thought our team had scored, so he jumped up and celebrated loudly. A few seconds later, we realised that the ball had actually missed the goal. Everyone around us burst out laughing, and my friend looked extremely embarrassed.",
            "Afterwards, we tried several rides and attended an outdoor music performance. Although we were exhausted by the end of the day, we felt relaxed and thoroughly enjoyed ourselves.",
            "Looking back, I realised that enjoyable experiences do not always need to be carefully planned. Sometimes, unexpected and even embarrassing moments become our best memories. Since then, I have become more willing to accept invitations and try activities outside my usual interests.",
        ],
        "focus": [
            "Sports match: kể kỹ trận đấu và khoảnh khắc ăn mừng hụt.",
            "Music festival: nhấn mạnh không khí và buổi biểu diễn.",
            "Amusement park: nhấn mạnh các trò chơi.",
            "Funny moment: tập trung vào phản ứng của người bạn.",
        ],
    },
    {
        "title": "Bài 6 — Chủ đề 11 + 12",
        "topics": (
            "Tall building / old building / work of art / museum / forest / "
            "extreme sport / outdoor challenge"
        ),
        "text": [
            "I would like to talk about a memorable trip to Ba Na Hills, where I experienced both impressive architecture and an outdoor challenge.",
            "I went there with several friends during our holiday in Da Nang. Before the trip, I had seen many photographs of the area, but I had never visited it myself. As the destination was located high in the mountains and surrounded by forest, we had to travel there by cable car.",
            "At first, I felt slightly frightened because the cable car was much higher than I had expected. However, the view of the forest and mountains was breathtaking, so I gradually became more relaxed.",
            "When we reached the top, we explored the French Village, which contained several tall buildings designed in an old European style. We also visited a small exhibition where we saw photographs, sculptures and other works of art. I was particularly impressed by a painting of the surrounding mountains because it captured the peaceful atmosphere perfectly.",
            "Later, my friends persuaded me to try an outdoor climbing activity. Although I was afraid of heights and nearly gave up, they encouraged me to continue. I eventually completed it and felt extremely proud of myself.",
            "Looking back, the trip was not only entertaining but also educational. It taught me that historic buildings and works of art can help us understand a place more deeply. I also realised that stepping outside our comfort zone can make us more confident.",
        ],
    },
    {
        "title": "Bài 7 — Chủ đề 14 + 15",
        "topics": (
            "Did something you did not want to do / broke a rule / someone asked you "
            "to stop / someone was rude to you / uncomfortable situation"
        ),
        "text": [
            "I would like to talk about an uncomfortable experience I had while preparing for a university presentation.",
            "It happened last year when my group had to give a presentation in front of the entire class. To be honest, I did not want to be the presenter because I was extremely shy and afraid of speaking in public. However, our group leader explained that everyone needed to take responsibility for one section, so I reluctantly agreed to present the introduction.",
            "The evening before the presentation, we decided to practise in an empty classroom. We knew that students were not supposed to use the room after closing time without permission. Nevertheless, because we were worried about the presentation, we stayed longer than we should have.",
            "After about thirty minutes, a security guard entered and told us to stop. Unfortunately, he spoke to us in an unnecessarily rude manner and raised his voice. I felt embarrassed and uncomfortable because several other students were nearby. Although I was annoyed, I remained calm, apologised and asked my teammates to leave without arguing.",
            "The following day, our presentation went considerably better than expected. I felt relieved and proud because I had overcome my fear of public speaking.",
            "Looking back, I realised that we had broken a rule even though we had not intended to cause trouble. The experience taught me to respect regulations and remain polite when someone behaves rudely. Since then, I have become better at controlling my reactions in uncomfortable situations.",
        ],
    },
    {
        "title": "Bài 8 — Chủ đề 16",
        "topics": "Sleeping habits / favourite outfit / personal routine",
        "text": [
            "I would like to talk about my evening routine and how it affects my sleeping habits.",
            "On weekdays, I normally try to go to bed at around ten thirty and wake up at six the following morning. I prefer going to bed early because getting enough sleep helps me remain energetic and productive throughout the day.",
            "Before going to bed, I usually take a shower, prepare my clothes for the following morning and listen to relaxing music. My favourite sleeping outfit is a loose T-shirt and a pair of comfortable shorts. I prefer soft, breathable clothes because they help me relax and sleep more comfortably.",
            "However, my routine is not always perfect. When I have too much work or spend too much time using my phone, I occasionally stay up much later than I should. If I do not get enough sleep, I tend to feel exhausted and find it difficult to concentrate the following morning.",
            "To improve this habit, I have started putting my phone away at least thirty minutes before bedtime. I also try to complete important tasks earlier instead of leaving them until late at night.",
            "Overall, I believe that a consistent personal routine can significantly improve our quality of life. Good-quality sleep is not only beneficial to our physical health but also essential for our mental well-being. Since adopting these habits, I have become healthier, happier and more productive.",
        ],
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(11)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.08

for style_name, size, color in [
    ("Title", 24, "17365D"),
    ("Heading 1", 17, "17365D"),
    ("Heading 2", 13, "2F5597"),
]:
    style = styles[style_name]
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

footer = section.footer.paragraphs[0]
add_page_number(footer)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("APTIS GENERAL SPEAKING PART 4")
run.bold = True
run.font.name = "Arial"
run.font.size = Pt(25)
run.font.color.rgb = RGBColor(23, 54, 93)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("8 bài gộp bao phủ 16 nhóm chủ đề")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(47, 85, 151)

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = note.add_run(
    "Cấu trúc tương đồng • Dễ ghi nhớ • Có ngữ pháp và từ vựng hướng band cao"
)
run.italic = True
run.font.size = Pt(11)

doc.add_paragraph()
doc.add_heading("Cách sử dụng tài liệu", level=1)
for item in [
    "Speaking Part 4: 1 phút chuẩn bị, sau đó nói 2 phút và trả lời đủ 3 câu hỏi.",
    "Không cần đọc nguyên bài theo kiểu học thuộc. Hãy giữ cốt truyện và thay đổi trọng tâm theo câu hỏi.",
    "Trong 1 phút chuẩn bị, ghi từ khoá cho ba phần: hoàn cảnh – cảm xúc/hành động – quan điểm/bài học.",
    "Luyện nói thành tiếng và bấm giờ; ưu tiên sự trôi chảy, phát âm rõ và nội dung đúng trọng tâm.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Bản đồ 16 nhóm chủ đề → 8 bài", level=1)
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
headers = table.rows[0].cells
headers[0].text = "Bài học"
headers[1].text = "Nhóm chủ đề bao phủ"
set_repeat_table_header(table.rows[0])
for cell in headers:
    set_cell_shading(cell, "D9EAF7")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for run in cell.paragraphs[0].runs:
        run.bold = True
for lesson in LESSONS:
    cells = table.add_row().cells
    cells[0].text = lesson["title"]
    cells[1].text = lesson["topics"]
    for cell in cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

doc.add_heading("Khung ngữ pháp dùng chung", level=1)
for sentence in [
    "I would like to talk about a time when...",
    "It happened about... when...",
    "Before that, I had...",
    "At that moment, I felt... because...",
    "However, instead of giving up, I decided to...",
    "Although..., I managed to...",
    "In the end,...",
    "Looking back, I realised that...",
    "The experience taught me that...",
    "Since then, I have become...",
]:
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(sentence)
    run.bold = True

for index, lesson in enumerate(LESSONS):
    doc.add_page_break()
    doc.add_heading(lesson["title"], level=1)

    topic_heading = doc.add_paragraph()
    run = topic_heading.add_run("Chủ đề bao phủ: ")
    run.bold = True
    run.font.color.rgb = RGBColor(47, 85, 151)
    run = topic_heading.add_run(lesson["topics"])
    run.italic = True

    for paragraph_text in lesson["text"]:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.add_run(paragraph_text)

    if lesson.get("focus"):
        doc.add_heading("Cách đổi trọng tâm theo đề", level=2)
        for item in lesson["focus"]:
            doc.add_paragraph(item, style="List Bullet")

doc.add_page_break()
doc.add_heading("Checklist luyện Speaking Part 4", level=1)
for item in [
    "Tôi đã trả lời đủ cả ba câu hỏi.",
    "Tôi nói đúng chủ đề, không dành thời gian mô tả bức ảnh.",
    "Bài nói có mở đầu, diễn biến, kết quả và bài học/quan điểm.",
    "Tôi đã dùng câu phức như although, when, because, which hoặc if.",
    "Tôi đã dùng nhiều từ nối, chẳng hạn however, in addition, in the end và looking back.",
    "Tôi nói gần đủ hai phút nhưng không cố nói quá nhanh.",
    "Phát âm rõ ràng và hạn chế khoảng dừng dài.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph()
closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = closing.add_run("Học cốt truyện – nhớ từ khoá – điều chỉnh theo đúng ba câu hỏi.")
run.bold = True
run.font.color.rgb = RGBColor(23, 54, 93)

doc.save(OUTPUT)
print(OUTPUT)
