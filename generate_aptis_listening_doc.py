from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT = "/workspace/Key_Aptis_Listening_Real_Tests.docx"
SOURCE = "https://fr.scribd.com/document/798350621/Key-Aptis-Listening-Real-Tests"

PART1_ITEMS = [
    ("1", "Furniture", "cái gì không phải đồ nguyên bản (original)?"),
    ("2", "Request a transfer", "lời khuyên dành cho những người thiếu động lực trong công việc?"),
    ("3", "6.30pm", "gặp nhau lúc mấy giờ?"),
    ("4", "At a new shopping mall", "cô ấy đi mua đồ ở đâu?"),
    ("5", "Going to the theatre", "cô ấy thích làm gì nhất lúc rảnh?"),
    ("6", "Football", "môn thể thao mà cô ấy giỏi nhất?"),
    ("7", "35 minutes", "cô ấy mất bao lâu đi xe đạp?"),
    ("8", "On Thursday morning", "họp vào lúc nào?"),
    ("9", "Three o'clock", "mẹ gặp con lúc mấy giờ?"),
    ("10", "Red", "nhà của Jack màu gì?"),
    ("11", "Play football", "anh ấy làm gì sau giờ làm?"),
    ("12", "Art", "con trai cô ấy thích học môn gì?"),
    ("13", "Two weeks", "đi du lịch Ấn Độ mấy tuần?"),
    ("14", "Thursday 13th", "buổi hẹn mới vào ngày bao nhiêu?"),
    ("15", "Iced tea", "người đàn ông uống nước gì?"),
    ("16", "In the afternoons", "cô ấy thường viết vào thời gian nào?"),
    ("17", "By train", "anh ấy đi bằng phương tiện gì?"),
    ("18", "Sees her family", "cô ấy làm gì vào thứ bảy?"),
    ("19", "Photography", "vợ anh ấy thích làm gì?"),
    ("20", "A large stone", "tại sao cô ấy trở thành nhà khoa học?"),
    ("21", "Wednesday afternoon", "khi nào được chơi đá bóng ở trường học?"),
    ("22", "2000", "thị trấn sẽ xây bao nhiêu tòa nhà?"),
    ("23", "Teacher", "nhà văn này làm nghề gì trước đó?"),
    ("24", "Having the meeting without him", "anh ấy yêu cầu điều chỉnh buổi họp như thế nào?"),
    ("25", "10.15", "buổi họp lúc mấy giờ?"),
    ("26", "Suggest a drink", "anh ấy gọi điện để làm gì?"),
    ("27", "Eggs", "mẹ gọi con gái nhờ mua gì?"),
    ("28", "10.00", "hẹn nhau lúc mấy giờ?"),
    ("29", "Stay late at the office", "cô ấy bận gì mà không đón con được?"),
    ("30", "Action film", "cô ấy khuyên nên xem phim gì?"),
    ("31", "The town hall", "anh ấy đang ở đâu?"),
    ("32", "By bus", "đi làm bằng phương tiện gì?"),
    ("33", "In the morning", "chuyên gia khuyên nên ăn hoa quả vào buổi nào?"),
    ("34", "To have some quiet time", "cô ấy dậy sớm làm gì?"),
    ("35", "He has to drive to work", "anh ấy học lái xe để làm gì?"),
    ("36", "Clothes", "anh ấy mua gì trực tiếp tại shop?"),
    ("37", "A university area", "đang mô tả khu vực nào?"),
    ("38", "Persuading his family", "vấn đề lớn nhất của anh ấy là gì?"),
    ("39", "1500 years", "thành phố này bao nhiêu năm tuổi?"),
    ("40", "On the first floor", "văn phòng ở tầng mấy?"),
    ("41", "250 pounds", "anh ấy trả bao nhiêu để mua máy tính?"),
    ("42", "The town hall", "anh ấy muốn đi đâu vào ngày mai?"),
    ("43", "Food", "2 người mang gì đi picnic?"),
    ("44", "10,000", "số lượng dân cư ở 1 vùng quê?"),
    ("45", "At the park", "2 người hẹn gặp nhau ở đâu?"),
    ("46", "Black", "mua áo màu gì?"),
    ("47", "To work in business", "anh ấy chọn nghề nghiệp gì?"),
    ("48", "On Saturday", "hạn nộp bài tập là thứ mấy?"),
    ("49", "Computer", "tham gia khóa học về cái gì?"),
    ("50", "The mountain scenes", "điều gì thu hút cô ấy nhất ở bộ phim?"),
    ("51", "The ending", "2 người cùng thích điều gì về bộ phim?"),
    ("52", "A large stone", "động lực gì khiến cô ấy nghiên cứu khoa học?"),
    ("53", "To say thank you", "anh ấy gọi điện để nói gì?"),
    ("54", "Over 300,000 copies", "bán được bao nhiêu bản copies?"),
    ("55", "Market place", "hẹn gặp ở đâu để đi bus về nhà?"),
    ("56", "Glasses", "anh ấy để quên cái gì?"),
    ("57", "15", "anh ấy phải thuyết trình bao nhiêu phút?"),
    ("58", "The girl's team", "cô ấy chụp ảnh cho ai?"),
    ("59", "Goes for a walk", "gia đình này làm gì hầu hết các cuối tuần?"),
    ("60", "Long and red", "cô ấy muốn cái váy nào?"),
    ("61", "Bathroom", "phòng nào cô ấy thích nhất?"),
    ("62", "Stayed at home", "cô ấy làm gì cuối tuần trước?"),
    ("63", "Kitchen", "phòng nào to nhất?"),
    ("64", "Cold and wet", "thời tiết thế nào?"),
    ("65", "Three", "ấn phím mấy để mua máy tính mới?"),
    ("66", "She walks", "cô ấy đi phương tiện gì đến trường?"),
    ("67", "9.15", "tàu chạy lúc mấy giờ? (để đi công tác)"),
    ("68", "Short", "chị gái cô ấy trông như thế nào?"),
    ("69", "Tea", "chị gái cô ấy uống gì?"),
    ("70", "Outside a shop", "2 vợ chồng hẹn gặp nhau ở đâu?"),
    ("71", "7 o'clock", "ăn tối lúc mấy giờ?"),
    ("72", "Speak at the conference", "thầy giáo muốn cô ấy làm gì?"),
    ("73", "History classes", "nhớ gì nhất hồi còn đi học?"),
    ("74", "The south", "cô ấy đi đâu cho kỳ nghỉ?"),
    ("75", "The words", "điều gì đặc biệt ở bài hát mới?"),
    ("76", "Fire from the countryside", "điều gì gây ô nhiễm môi trường?"),
    ("77", "3250 pounds", "ô tô bé nhất giá bao nhiêu?"),
    ("78", "Room 301", "lớp học ở phòng nào?"),
    ("79", "21", "Stephanie bao nhiêu tuổi?"),
    ("80", "Sick", "anh ấy cảm thấy thế nào?"),
    ("81", "9.30", "tàu chạy lúc mấy giờ? (tàu bị delay)"),
    ("82", "Make plans later", "2 người hẹn nhau làm gì?"),
    ("83", "The mountains", "cô ấy sẽ đi đâu?"),
    ("84", "By the hotel's main entrance", "hẹn nhau ở đâu để đi xe bus?"),
    ("85", "The river boat", "trà được phục vụ ở đâu?"),
    ("86", "The city's favorite group", "ở cuối buổi hoà nhạc có điều gì?"),
    ("87", "20 10 30", "số điện thoại của cửa hàng?"),
    ("88", "Phone", "cô ấy làm mất cái gì?"),
    ("89", "Quarter to eight", "2 người hẹn nhau lúc mấy giờ?"),
    ("90", "Camping", "năm ngoái họ đi đâu?"),
    ("91", "One pound fifty", "dụng cụ vệ sinh giá bao nhiêu?"),
    ("92", "Where to buy a new table", "họ cần mua gì cho phòng khách?"),
    ("93", "In the east", "thời tiết ở đâu đẹp nhất?"),
    ("94", "Poor weather conditions", "tại sao chuyến bay bị huỷ?"),
    ("95", "Two", "đi cửa nào để đi tàu đến Edinburgh?"),
    ("96", "9am on Sunday", "2 người bạn gặp nhau lúc nào?"),
    ("97", "20 minutes", "mất bao lâu để đi đến ga?"),
    ("98", "Not enough people", "tại sao buổi thăm bảo tàng bị huỷ?"),
    ("99", "Cycling", "anh ấy hay làm gì năm ngoái?"),
    ("100", "Opposite the hotel", "văn phòng ở đâu?"),
    ("101", "Practical", "quan điểm của anh ấy về đi tàu?"),
    ("102", "Opposite the gift shop", "quán cafe ở đâu?"),
    ("103", "His sister and her children", "ai sắp đến thăm anh ấy?"),
    ("104", "Best friends", "cô ấy ở chung với ai?"),
    ("105", "Help people", "tại sao cô ấy muốn làm nhà văn?"),
    ("106", "A park", "câu lạc bộ ở gần đâu?"),
    ("107", "Fish", "cho mèo ăn gì?"),
    ("108", "Chocolates", "mua gì cho chị gái?"),
    ("109", "1 p.m", "hỏi về thời gian có football match"),
    ("110", "He taught her a lot", "tại sao cô ấy thích người quản lý?"),
    ("111", "Writer", "anh ấy muốn trở thành gì?"),
    ("112", "Water", "anh ấy uống gì?"),
    ("113", "Go for a drive", "cô ấy sẽ làm gì?"),
    ("114", "They have similar interests", "cô ấy giống mẹ ở đâu?"),
    ("115", "A suit for the office", "mua gì ở cửa hàng?"),
    ("116", "They stay in groups for protection", "chim làm gì vào mùa đông?"),
    ("117", "1.50", "trứng bao nhiêu tiền?"),
    ("118", "2pm", "mấy giờ có bữa trưa?"),
    ("119", "2pm", "mấy giờ họp?"),
    ("120", "Go for a walk", "cô ấy thường làm gì vào buổi tối?"),
    ("121", "Tuesday", "thứ mấy gặp nhau?"),
    ("122", "She works irregular time", "làm nhà văn thì có gì khác với các nghề khác?"),
    ("123", "At a new shopping centre", "mua đồ ăn ở đâu?"),
    ("124", "1pm", "trận bóng đá lúc mấy giờ?"),
    ("125", "Tea", "cô ấy uống gì cho bữa trưa?"),
    ("126", "White", "nhà của giáo viên màu gì?"),
    ("127", "2.50 pounds", "trứng bao nhiêu tiền?"),
    ("128", "Walking", "cô ấy làm gì vào kỳ nghỉ?"),
    ("129", "Clothes", "2 người cùng mua gì?"),
    ("130", "Speaking at the conference", "cô ấy phải làm gì?"),
    ("131", "A university area", "đang nói về khu vực nào?"),
    ("132", "They were thin", "mẹ và dì giống nhau ở điểm nào?"),
    ("133", "One", "có bao nhiêu người Mỹ?"),
    ("134", "2 weeks", "đi du lịch Ấn Độ mấy tuần?"),
    ("135", "Go to the theatre and play sports", "cô ấy thường làm gì?"),
    ("136", "Go to the park", "đi đâu khi đi du lịch?"),
    ("137", "Play golf", "có hoạt động gì vào buổi chiều?"),
    ("138", "The college", "đi bộ đi đâu mỗi tối?"),
    ("139", "20", "chuẩn bị bao nhiêu ghế cho cuộc họp?"),
    ("140", "In the corner", "gọi điện cho quán cà phê nhờ tìm đồ để quên ở đâu?"),
    ("141", "The front entrance", "tìm thấy chìa khóa ở đâu?"),
    ("142", "A dress", "con gái gọi điện cho bố đã mua gì?"),
    ("143", "Order the food", "sắp có school party, giáo viên phải chuẩn bị gì?"),
    ("144", "Friday", "cần máy tính khi nào?"),
    ("145", "France", "kỳ tới học về đất nước nào?"),
    ("146", "Windows", "sửa chữa gì cho building?"),
    ("147", "Drawing", "diễn viên thích làm gì?"),
    ("148", "A Performance space", "trường học sắp xây gì mới?"),
]

PART2_TOPICS = [
    ("Protect the environment", [
        ("A", "Buy environmentally friendly products"),
        ("B", "Give away used items"),
        ("C", "Does not use commercial cleaning products"),
        ("D", "Reuse containers for storing food"),
    ]),
    ("Online shopping", [
        ("A", "The products are delivered"),
        ("B", "There are more choices"),
        ("C", "It saves time"),
        ("D", "It is cheaper"),
    ]),
    ("Studying", [
        ("A", "In various places"),
        ("B", "In a quiet place"),
        ("C", "With music"),
        ("D", "Late at night"),
    ]),
    ("Studying (2)", [
        ("A", "In the coffee shop"),
        ("B", "In the park"),
        ("C", "On public transport"),
        ("D", "At home"),
    ]),
    ("Listening to music", [
        ("A", "To relax"),
        ("B", "While studying"),
        ("C", "While singing"),
        ("D", "After waking up"),
    ]),
    ("Protecting the environment (2)", [
        ("A", "Using less electricity"),
        ("B", "Not driving to work"),
        ("C", "Shopping online"),
        ("D", "Using less water"),
    ]),
    ("Outdoor activities", [
        ("A", "Horse riding"),
        ("B", "Going for a run"),
        ("C", "Climbing"),
        ("D", "Mountain biking"),
    ]),
    ("The place to run", [
        ("A", "In the fitness centre"),
        ("B", "At the seaside"),
        ("C", "In the street"),
        ("D", "On the running track"),
    ]),
    ("Do exercise", [
        ("A", "Hate exercising"),
        ("B", "Help improve work performance"),
        ("C", "Find exercise tiring"),
        ("D", "Has fun when exercising with friends"),
    ]),
    ("Use the internet", [
        ("A", "Watch films"),
        ("B", "Communicate with friends"),
        ("C", "Complete assignments"),
        ("D", "Find transport information"),
    ]),
    ("Art", [
        ("A", "Social activity"),
        ("B", "With children"),
        ("C", "As part of their job"),
        ("D", "Alone"),
    ]),
    ("Travel to work", [
        ("A", "Travel by bus"),
        ("B", "Travel by car"),
        ("C", "Walk alone"),
        ("D", "Walk with a friend"),
    ]),
]

PART3_TOPICS = [
    ("Changes in the workplace (Man start)", [
        ("Continuity is important when planning a career", "Man"),
        ("Job security cannot be guaranteed", "Woman"),
        ("Job satisfaction is important for motivator", "Both"),
        ("Technological improvement is good for the economy", "Man"),
    ]),
    ("Changes in the workplace (Woman start)", [
        ("Continuity is important when planning a career", "Woman"),
        ("Job security cannot be guaranteed", "Man"),
        ("Job satisfaction is important for motivator", "Both"),
        ("Technological improvement is good for the economy", "Woman"),
    ]),
    ("The subject of beauty (MWBW)", [
        ("People share the very similar ideas about beauty", "Man"),
        ("Views of beauty change over time", "Woman"),
        ("Beauty can be found in unlikely places", "Both"),
        ("Traditional ideas of beauty are going to change", "Woman"),
    ]),
    ("The subject of beauty (variant 2)", [
        ("People share the very similar ideas about beauty", "Woman"),
        ("Views of beauty change over time", "Man"),
        ("Beauty can be found in unlikely places", "Both"),
        ("Traditional ideas of beauty are going to change", "Man"),
    ]),
    ("Actors (MWBB)", [
        ("Auditions are the most important of casting", "Man"),
        ("Actors respond best to a strong script", "Woman"),
        ("Theatre acting and movie acting require different skills", "Both"),
        ("Actors need to be praised", "Both"),
    ]),
    ("Information and technology (MWWB)", [
        ("The future generation will fail to cope with new information", "Man"),
        ("The information revolution will be good for the economy", "Woman"),
        ("No computer is superior to the human brain", "Woman"),
        ("More should be done to protect individual privacy", "Both"),
    ]),
    ("Internet (BWWB)", [
        ("There is too much information on the internet", "Both"),
        ("Using internet requires skills", "Woman"),
        ("The internet is changing the way we think", "Woman"),
        ("Internet has made people less patient", "Both"),
    ]),
    ("Arts (WMBM)", [
        ("Art is only suitable for the privileged few", "Woman"),
        ("The government should invest more in arts", "Man"),
        ("Children should be exposed to art early", "Both"),
        ("In the future, art will become more accessible", "Man"),
    ]),
    ("Music (MBWB)", [
        ("Singers play a good role for young people", "Man"),
        ("Taste in music is a highly personal thing", "Both"),
        ("Music is a universal language", "Woman"),
        ("Music can manipulate people's feelings", "Both"),
    ]),
    ("University and technology (BMWW)", [
        ("Technology helps make education more accessible", "Both"),
        ("Social interaction is important", "Man"),
        ("The diverse curriculum is not an advantage", "Woman"),
        ("University competition should be encouraged", "Woman"),
    ]),
    ("Politics (BMWM)", [
        ("Young people are becoming more interested in politics", "Both"),
        ("Social media has changed political activism", "Man"),
        ("People are better informed political issues", "Woman"),
        ("More women are likely to participate in politics", "Man"),
    ]),
    ("Urban farming (WMMB)", [
        ("Living space is more important than farming space", "Woman"),
        ("Urban farming sites can be visually appealing", "Man"),
        ("Urban farming can benefit local economics", "Man"),
        ("Urban farming cannot meet food needs", "Both"),
    ]),
    ("Community (BWMM)", [
        ("Building design can influence people's behavior", "Both"),
        ("Creating community can take time", "Woman"),
        ("Work communities and social communities are the same", "Man"),
        ("Technology has changed how community forms", "Man"),
    ]),
    ("Information technology (MWWB - đề mới)", [
        ("Future generations fail to cope with technology information", "Man"),
        ("Technology revolution is good for the economy", "Woman"),
        ("No computer is superior to the human brain", "Woman"),
        ("More should be done to protect individual privacy", "Both"),
    ]),
]

PART4_TOPICS = [
    ("A break from studying", [
        "He wasn't ready to start higher education",
        "To be more independent",
    ]),
    ("A book about a life of a scientist", [
        "It uses simple language to describe complex ideas",
        "It is similar to the previous book about the scientist",
    ]),
    ("Sport", [
        "They have a harmful effect",
        "It helps balance student's lives",
    ]),
    ("Television series", [
        "It has the consistent quality throughout",
        "Viewer habits influence the way that series are made",
    ]),
    ("Advertising and sponsoring", [
        "It can help reach new customers",
        "They can generate negative publicity for the sport",
    ]),
    ("2 famous writers", [
        "They both make references to each other in their work",
        "The meaning of their work is not always easily identified",
    ]),
    ("A regional development plan", [
        "It doesn't do enough to promote alternatives to driving",
        "The plan is not making efficient use of existing land",
    ]),
    ("Life after university", [
        "Be flexible and open minded",
        "They are becoming more competitive",
    ]),
    ("Security cameras", [
        "Employees probably worry unnecessarily",
        "People should feel reassured",
    ]),
    ("A new novel of a famous writer", [
        "It is quite different compared to his previous works",
        "He should listen to the critics before writing the next novel",
    ]),
    ("A musician's life", [
        "He will probably retire from singing",
        "He could have been more successful",
    ]),
    ("A new guide", [
        "It creates a sense of adventure",
        "It is only suitable for a particular generation",
    ]),
    ("A research into happiness", [
        "It has not been accurately reported by the media",
        "The research is unlikely to find a conclusive answer",
    ]),
    ("Criticism of the new novel", [
        "The characters were interesting",
        "It will establish the author's popularity",
    ]),
    ("Writer's block", [
        "Create dedicated periods",
        "Refusing to seek the advice of other people",
    ]),
    ("Professionalism at work", [
        "Maintain the positive attitude",
        "Change its definition",
    ]),
    ("Making plans", [
        "They help you make decisions",
        "It requires you to set yourself certain limits",
    ]),
    ("A promotion campaign for a product", [
        "It is making exaggerated claims",
        "It costs too much to make to be profitable",
    ]),
    ("Script production", [
        "The characters' backgrounds are not adequately explored",
        "The new industry demands are negatively influencing script production",
    ]),
    ("A new restaurant", [
        "The standard of service wasn't good",
        "They need to make the customers feel valued and welcome",
    ]),
    ("Work from home", [
        "It wasn't as good as she expected",
        "It depends on your situation and personality",
    ]),
    ("Managing financial spending", [
        "Monitor your spending for a weekly plan",
        "Seek advice from someone who has experience",
    ]),
    ("The importance of sleep", [
        "Blocking out noise and light is key",
        "People can't always recognize the symptoms of tiredness",
    ]),
    ("A newly broadcasted TV series (đề mới 01/2025)", [
        "It caught the audience's attention from the start",
        "It has also helped its network reach new customers",
    ]),
    ("The advertising industry (dạng 1 - đề mới)", [
        "Series are damaged by overexposure",
        "They can generate negative publicity for the sport",
    ]),
    ("The advertising industry (dạng 2 - đề mới)", [
        "It helps to reach new customers",
        "They are not always good for sport fans",
    ]),
    ("A new series (đề mới)", [
        "The dialogues seem unrealistic",
        "The new industry demand is negatively influencing script production",
    ]),
    ("Promotion campaign for a product (dạng 2 - đề mới)", [
        "It is not targeting the correct market",
        "It is too similar to many existing products",
    ]),
    ("Manage personal finances (đề mới)", [
        "Organizing their resources more effectively",
        "Get advice from people that have experience",
    ]),
    ("Security cameras (dạng 2 - đề mới)", [
        "People are unnecessarily worried about them",
        "People should feel reassured",
    ]),
    ("A book about a life of a scientist (dạng 2 - đề mới)", [
        "It is exciting to read",
        "It has been written for a general audience",
    ]),
    ("Writer's new novel (đề mới)", [
        "It is different from his earlier works",
        "He should listen to critics before writing his next work",
    ]),
    ("Goal setting (đề mới)", [
        "It allows you to be more flexible",
        "It can prevent you from making mistakes",
    ]),
    ("The latest scripts for a new show (đề mới)", [
        "They seem unrealistic",
        "It is negatively influencing script production",
    ]),
]


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size in [("Title", 24), ("Heading 1", 16), ("Heading 2", 13)]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(23, 54, 93)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    add_page_number(section.footer.paragraphs[0])


def add_title_block(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Key Aptis Listening Real Tests")
    run.bold = True
    run.font.size = Pt(24)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Hệ thống trọng điểm nghe — Đầy đủ Part 1–4")
    run.bold = True
    run.font.size = Pt(14)

    for line in [
        "APTIS REAL TESTS",
        "HỆ THỐNG TRỌNG ĐIỂM NGHE (UPDATE LIÊN TỤC)",
        f"Nguồn: {SOURCE}",
        "NOTE: Thí sinh chú ý chủ động vào link này check thường xuyên. Phần update sẽ để ở đầu các part để các bạn dễ nhận thấy.",
    ]:
        p = doc.add_paragraph(line)
        if line.startswith("NOTE"):
            p.runs[0].italic = True


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    doc.add_paragraph()


doc = Document()
setup_styles(doc)
add_title_block(doc)

doc.add_heading("Tổng quan cấu trúc", level=1)
for item in [
    "PART 1: 13 câu nghe ngắn — chọn 1 trong 3 đáp án (key dưới đây là đáp án đúng).",
    "PART 2: Câu 14 — 4 người A/B/C/D nói về cùng 1 chủ đề, ghép với 6 ý.",
    "PART 3: Câu 15 — hội thoại nam/nữ, chọn Man / Woman / Both cho 4 phát biểu.",
    "PART 4: Câu 16–17 — 2 bài nói, mỗi bài 2 câu hỏi (đề thật nhặt 2 trong số các topic).",
    "Lưu ý: Hội đồng Anh hay tráo đề, tráo thứ tự người nói và tráo giọng nam/nữ. Không học tủ 1 phiên bản.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_page_break()
doc.add_heading("PART 1", level=1)
doc.add_paragraph("Note: Trong đề thi thật sẽ nhặt ra 13 câu trong số các câu bên dưới:")
add_table(doc, ["STT", "Key / Đáp án", "Câu hỏi (tiếng Việt)"], PART1_ITEMS)

doc.add_page_break()
doc.add_heading("PART 2", level=1)
doc.add_paragraph(
    "Note: Câu 14 có 4 người A, B, C, D cùng nói về một chủ đề. "
    "Hội đồng Anh hay tráo thứ tự người nói — phải nghe kỹ để ghép đúng."
)
for topic, options in PART2_TOPICS:
    doc.add_heading(f"Topic: {topic}", level=2)
    add_table(doc, ["Người nói", "Key / Ý chính"], options)

doc.add_page_break()
doc.add_heading("PART 3", level=1)
doc.add_paragraph(
    "Note: Hội đồng Anh rất hay tráo giọng nói nam - nữ, dẫn đến đáp án Woman/Man bị đảo vị trí. "
    "Đáp án Both thường giữ nguyên. Phải nghe cẩn thận."
)
for topic, statements in PART3_TOPICS:
    doc.add_heading(f"Topic: {topic}", level=2)
    add_table(doc, ["Phát biểu", "Đáp án"], statements)

doc.add_page_break()
doc.add_heading("PART 4", level=1)
doc.add_paragraph(
    "Note: Đề thi thật nhặt 2 trong số các topic dưới đây. "
    "Câu 16 và 17 mỗi câu có 2 đáp án đúng."
)
for topic, answers in PART4_TOPICS:
    doc.add_heading(f"Topic: {topic}", level=2)
    for idx, answer in enumerate(answers, start=1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(answer)

doc.save(OUTPUT)
print(OUTPUT)
print(f"PART1: {len(PART1_ITEMS)} | PART2 topics: {len(PART2_TOPICS)} | PART3 topics: {len(PART3_TOPICS)} | PART4 topics: {len(PART4_TOPICS)}")
